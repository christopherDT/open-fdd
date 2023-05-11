import pandas as pd
from pydantic import BaseModel
from typing import Optional
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import math
import os
import time
from io import BytesIO
from num2words import num2words

class Sensor(BaseModel):
    """Gives the different attributes for relevant sensors used in each fault.

    Atrributes:
        col_name            The standard column name any time this sensor is in a Pandas data frame
        long_name           The standard full sensor name, how it's referred to in fault reports
        measurement         How the sensor is measured (temperature, percent, pressure, etc.)
        short_name          The sensor's corresponding short name, mainly used for legends in graphs
    """
    col_name: str
    long_name: str
    measurement: str
    units: str
    short_name: Optional[str] = None
    avg_fault_val: Optional[float] = None

    def __init__(self,**data):
        super().__init__(**data)
        if 'short_name' not in data.keys():
            super().__init__(short_name = data['long_name'].title(),**data)

class Fault(BaseModel):
    """
    Gives a fault's relevant attributes for generating final reports.

    Attributes:
        num         The fault number as an integer
        definition  The text description for the fault's definition
        sensors     All of the fault's relevant sensors (not including supply_vfd_speed, which is always assumed to exist)

    """
    num: int
    definition: str
    sensors: list

    def __init__(self, col_names, **data):
        # do this so we can get the sensors from ALL_SENSORS whose col_names match the provided col_name
        super().__init__(sensors = [sensor for sensor in ALL_SENSORS if sensor.col_name in col_names],**data)

class ReportCalculator:
    """Calculates the data fed to the report. Assumes 'supply_vfd_speed' col exists."""
    def __init__(self, fault, df):
        self.fault = fault

        output_col = f"fc{fault.num}_flag"

        delta = df.index.to_series().diff()

        self.total_days = round(delta.sum() / pd.Timedelta(days=1), 2)

        self.total_hours = delta.sum() / pd.Timedelta(hours=1)

        self.hours_fault_mode = (delta * df[output_col]).sum() / pd.Timedelta(hours=1)

        self.percent_in_fault_mode = round(df[output_col].mean() * 100, 2)

        # get average fault values for each sensor
        for sensor in self.fault.sensors:
            sensor.avg_fault_val = round(df[sensor.col_name].where(df[output_col] == 1).mean(), 2)

        motor_on = df['supply_vfd_speed'].gt(1.).astype(int)
        self.hours_motor_runtime = round(
            (delta * motor_on).sum() / pd.Timedelta(hours=1), 2)

        # for summary stats on I/O data to make useful
        self.df_motor_on_filtered = df[df['supply_vfd_speed'] > 1.0]
        
class DocumentGenerator:
    """Class provides the skeleton for creating a report document."""

    def __init__(self, fault, df):
        self.fault = fault
        self.df = df
        self.output_col = f'fc{self.fault.num}_flag'
        self.document = Document()


    # creates combination timeseries and fault flag plots 
    def create_dataset_plot(self) -> plt:
        # group key sensors by measurement type
        import itertools
        grouped_sensors = []
        for key, group in itertools.groupby(self.fault.sensors, key=lambda x: x.measurement):    
            grouped_sensors.append(list(group))

        n_groups = len(grouped_sensors)

        fig, (*data_axes, fault_flag_ax) = plt.subplots(n_groups + 1, 1, figsize=(25, 8))
        plt.title(f'Fault Conditions {self.fault.num} Plot')

        # data_axes is the different axes grouped by measurement, so they can be graphed on similar scales
        for i in range(n_groups):
            group = grouped_sensors[i]
            for sensor in group:
                data_axes[i].plot(self.df.index, self.df[sensor.col_name], label=sensor.short_name)

            measurement = list(set([sensor.measurement for sensors in group]))[0]
            units = list(set([sensor.units for sensors in group]))[0]

            data_axes[i].legend(loc='best')
            data_axes[i].set_ylabel(f'{measurement} ({units})')

        # fault flag plot
        fault_flag_ax.plot(self.df.index, self.df[f'fc{self.fault.num}_flag'], label="Fault", color="k")
        fault_flag_ax.set_xlabel('Date')
        fault_flag_ax.set_ylabel('Fault Flags')
        fault_flag_ax.legend(loc='best')

        plt.legend()
        plt.tight_layout()

        return fig

    # creates a histogram plot for times when fault condition is true
    def create_hist_plot(self) -> plt:
        # calculate dataset statistics
        self.df[f"hour_of_the_day_fc{self.fault.num}"] = self.df.index.hour.where(self.df[self.output_col] == 1)

        # make hist plots fc10
        fig, ax = plt.subplots(tight_layout=True, figsize=(25, 8))
        ax.hist(self.df[f"hour_of_the_day_fc{self.fault.num}"].dropna())
        ax.set_xlabel("24 Hour Number in Day")
        ax.set_ylabel("Frequency")
        ax.set_title(f"Hour-Of-Day When Fault Flag {self.fault.num} is TRUE")
        return fig

    def create_report(self,path: str) -> None:
        print(f"Starting {path} docx report!")

        # add fault definition
        self.document.add_heading(f"Fault Condition {num2words(self.fault.num).title()} Report", 0)

        p = self.document.add_paragraph(self.fault.definition)

        self.document.add_picture(
            os.path.join(os.path.curdir, "images", f"fc{self.fault.num}_definition.png"),
            width=Inches(6),
        )

        # add dataset plot
        self.document.add_heading("Dataset Plot", level=2)

        fig = self.create_dataset_plot()
        fan_plot_image = BytesIO()
        fig.savefig(fan_plot_image, format="png")
        fan_plot_image.seek(0)

        # ADD IN SUBPLOTS SECTION
        self.document.add_picture(
            fan_plot_image,
            width=Inches(6),
        )

        # add dataset statistics
        self.document.add_heading("Dataset Statistics", level=2)

        calculator = ReportCalculator(self.fault, self.df)
    
        stats_lines = [
            f"Total time in days calculated in dataset: {calculator.total_days}",
            f"Total time in hours calculated in dataset: {calculator.total_hours}",
            f"Total time in hours for when fault flag is True: {calculator.hours_fault_mode}",
            f"Percent of time in the dataset when the fault flag is True: {calculator.percent_in_fault_mode}%",
            f"Percent of time in the dataset when the fault flag is False: {round((100 - calculator.percent_in_fault_mode), 2)}%",
            f"Calculated motor runtime in hours based off of VFD signal > zero: {calculator.hours_motor_runtime}"
        ]

        for line in stats_lines:
            paragraph = self.document.add_paragraph()
            paragraph.style = "List Bullet"
            paragraph.add_run(line)

        paragraph = self.document.add_paragraph()

        # if there are faults, add the histogram plot
        fc_max_faults_found = self.df[self.output_col].max()
        if fc_max_faults_found != 0:

            self.document.add_heading("Time-of-day Histogram Plots", level=2)
            histogram_plot_image = BytesIO()
            histogram_plot = self.create_hist_plot()
            histogram_plot.savefig(histogram_plot_image, format="png")
            histogram_plot_image.seek(0)
            self.document.add_picture(
                histogram_plot_image,
                width=Inches(6),
            )

            paragraph = self.document.add_paragraph()

            max_faults_line = f'When fault condition {self.fault.num} is True, the'

            for sensor in calculator.fault.sensors:
                if sensor.measurement == 'temperature': # only temp sensors
                    max_faults_line =  f'{max_faults_line} average {sensor.long_name} is {sensor.avg_fault_val} {sensor.units}, ' # need to add {units}

            max_faults_line = max_faults_line[0:-2] + "."

        else:
            print("NO FAULTS FOUND - For report skipping time-of-day Histogram plot")
            max_faults_line = f'No faults were found in this given dataset for the equation defined by ASHRAE.'


        paragraph.style = 'List Bullet'
        paragraph.add_run(max_faults_line)

        # ADD in Summary Statistics
        self.document.add_heading(
            'Summary Statistics filtered for when the AHU is running', level=1)

        for sensor in calculator.fault.sensors:
            if sensor.measurement == 'temperature': # only temp sensors
                self.document.add_heading(sensor.short_name, level=3)
                paragraph = self.document.add_paragraph()
                paragraph.style = 'List Bullet'
                paragraph.add_run(str(calculator.df_motor_on_filtered[sensor.col_name].describe()))

        self.document.add_heading("Suggestions based on data analysis", level=3)
        paragraph = self.document.add_paragraph()
        paragraph.style = "List Bullet"

        if calculator.percent_in_fault_mode > 5.0:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is high indicating temperature sensor error or the cooling valve is stuck open or leaking causing overcooling. Trouble shoot a leaking valve by isolating the coil with manual shutoff valves and verify a change in AHU discharge air temperature with the AHU running.')

        else:
            paragraph.add_run(
                'The percent True metric that represents the amount of time for when the fault flag is True is low inidicating the AHU components are within calibration for this fault equation Ok.')

        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f"Report generated: {time.ctime()}")
        run.style = "Emphasis"
        return self.document


sensor_attrs = ['col_name','long_name','measurement', 'units', 'short_name']

# this is just a simple way to store the sensors before we make them into objects
# this is ordered as: col_name, long_name, measurement, short_name, where short_name is optional
# this should probably not live in code, would be better as a csv or json or database or whatever
slist = [['mat', 'mixing air temperature', 'temperature', '°F', 'Mix Temp'],
        ['rat', 'return air temperature', 'temperature', '°F', 'Return Temp'],
        ['oat', 'outside air temperature', 'temperature', '°F', 'Out Temp'],
        ['satsp', 'supply air temperature setpoint', 'temperature', '°F', 'Out Temp',],
        ['economizer_sig', 'outside air damper position', 'percent', '%', 'AHU Dpr Cmd'],
        ['clg', 'AHU cooling valve', 'percent', '%', 'AHU Cool Valv'],
        ['htg', 'AHU heating valve', 'percent', '%', 'AHU Htg Valv'],
        ['cooling_sig', 'cooling signal', 'operating state', '', 'Mech Clg'],
        ['heating_sig', 'heating signal', 'operating state', '', 'Heat'],
        ['vav_total_flow', 'VAV total air flow', 'volume flow', 'cfm', 'VAV total flow'],
        ['duct_static', 'duct static pressure', 'pressure','inches WC', 'static']

        # ['supply_vfd_speed', 'supply fan speed', 'speed']
    ]

ALL_SENSORS = [Sensor(**dict(zip(sensor_attrs, sensor))) for sensor in slist]


fault_attrs = ['num', 'col_names', 'definition']

# same thing here -- this shouldn't be stored as code, should be separate files that can be easily tweaked
fault_defs = [
    [1, ['duct_static', 'duct_static_setpoint', 'supply_vfd_speed'], """Fault condition one of ASHRAE Guideline 36 is related to flagging poor performance of a AHU variable supply fan attempting to control to a duct pressure setpoint. Fault condition equation as defined by ASHRAE:"""],
    [2, ['mat', 'rat', 'oat'], """Fault condition two and three of ASHRAE Guideline 36 is related to flagging mixing air temperatures of the AHU that are out of acceptable ranges. Fault condition 2 flags mixing air temperatures that are too low and fault condition 3 flags mixing temperatures that are too high when in comparision to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges. Fault condition two equation as defined by ASHRAE:"""],
    [3, ['mat','rat','oat'], """Fault condition two and three of ASHRAE Guideline 36 is related to flagging mixing air temperatures of the AHU that are out of acceptable ranges. Fault condition 2 flags mixing air temperatures that are too low and fault condition 3 flags mixing temperatures that are too high when in comparision to return and outside air data. The mixing air temperatures in theory should always be in between the return and outside air temperatures ranges. Fault condition three equation as defined by ASHRAE:"""],
    [4, ['economizer_sig','heating_sig','cooling_sig'], """Fault condition four of ASHRAE Guideline 36 is related to flagging AHU control programming that is hunting between heating, economizing, economizing plus mechanical cooling, and mechanical cooling operating states. This fault diagnostic does NOT flag simultaneous heating and cooling, just excessive cycling between the states or operating modes the AHU maybe going in and out of. Fault condition four equation as defined by ASHRAE:"""],
    [9, ['satsp', 'oat'], """Fault condition nine of ASHRAE Guideline 36 is an AHU economizer free cooling mode only with an attempt at flagging conditions where the outside air temperature is too warm for cooling without additional mechanical cooling. Fault condition nine equation as defined by ASHRAE:"""],
    [10, ['oat', 'mat', 'cooling_sig', 'economizer_sig'], """Fault condition ten of ASHRAE Guideline 36 is an AHU economizer + mechanical cooling mode only with an attempt at flagging conditions where the outside air temperature and mixing air temperatures are not approximetely equal when the AHU is in a 100% outside air mode. Fault condition ten equation as defined by ASHRAE:"""]
]

ALL_FAULTS = [Fault(**dict(zip(fault_attrs,fault_def))) for fault_def in fault_defs]

FAULT_NUM = 4

# -------------------- FC 10
# OUTDOOR_DEGF_ERR_THRES = 5.
# MIX_DEGF_ERR_THRES = 5.
# RETURN_DEGF_ERR_THRES = 2.

# from faults import FaultConditionTwo
# _fc2 = FaultConditionTwo(
#     OUTDOOR_DEGF_ERR_THRES,
#     MIX_DEGF_ERR_THRES,
#     RETURN_DEGF_ERR_THRES,
#     "mat",
#     "rat",
#     "oat",
#     "supply_vfd_speed"
# )

# #-------------------- FC 10
# # ADJUST this param for the AHU MIN OA damper stp
# AHU_MIN_OA = 20

# # G36 params shouldnt need adjusting
# # error threshold parameters
# OAT_DEGF_ERR_THRES = 5
# MAT_DEGF_ERR_THRES = 5

# from faults import FaultConditionTen
# FAULT_CONDITION = FaultConditionTen(
#     OAT_DEGF_ERR_THRES,
#     MAT_DEGF_ERR_THRES,
#     "mat",
#     "oat",
#     "clg",
#     "economizer_sig",
# )

# -------------------- FC 4
# G36 params COULD need adjusting
# default OS MAX state changes is 7
# which seems high, could be worth adjusting 
# down to 4 to see what the faults look like
DELTA_OS_MAX = 7

# ADJUST this param for the AHU MIN OA damper stp
AHU_MIN_OA = 20

from faults import FaultConditionFour
FAULT_CONDITION = FaultConditionFour(
    DELTA_OS_MAX,
    AHU_MIN_OA,
    "economizer_sig",
    "heating_sig",
    "cooling_sig",
    "supply_vfd_speed"
)

DF = pd.read_csv(f"ahu_data/hvac_random_fake_data/fc{FAULT_NUM}_fake_data1.csv", index_col="Date", parse_dates=True).rolling("5T").mean()

df2 = FAULT_CONDITION.apply(DF)

FAULT_OBJ = [fault for fault in ALL_FAULTS if fault.num == FAULT_NUM][0]

# breakpoint()

# calculator = ReportCalculator(FAULT_OBJ, df2)

document_generator = DocumentGenerator(FAULT_OBJ, df2)

# breakpoint()
# print(calculator.summarize_fault_times())

# dataset_fig = document_generator.create_dataset_plot()
# hist_fig = document_generator.create_hist_plot()
# fig = document_generator_generator.create_dataset_plot()
# plt.show()

report = document_generator.create_report('blah')
report.save(f"test_fc{FAULT_NUM}.docx")